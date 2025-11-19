"""
Модуль для парсинга документов (PDF и DOCX) с поддержкой различных кодировок
"""

import os
import io
import chardet
import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional, Tuple
import config


class DocumentParser:
    """Парсер для документов PDF и DOCX"""
    
    def __init__(self):
        self.supported_formats = config.SUPPORTED_FORMATS
        self.supported_encodings = config.SUPPORTED_ENCODINGS
    
    def parse(self, file_path: str) -> Tuple[str, dict]:
        """
        Парсит документ и возвращает текст
        
        Args:
            file_path: путь к файлу
            
        Returns:
            Tuple[str, dict]: (текст документа, метаданные)
        """
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            return self.parse_pdf(file_path)
        elif extension == '.docx':
            return self.parse_docx(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {extension}")
    
    def parse_pdf(self, file_path: str) -> Tuple[str, dict]:
        """
        Парсит PDF документ
        
        Args:
            file_path: путь к PDF файлу
            
        Returns:
            Tuple[str, dict]: (текст документа, метаданные)
        """
        text_parts = []
        metadata = {
            'pages': 0,
            'format': 'pdf',
            'encoding': None
        }
        
        try:
            # Используем pdfplumber для лучшего извлечения текста
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                if metadata['pages'] > config.MAX_PAGES:
                    raise ValueError(f"Документ содержит {metadata['pages']} страниц, "
                                   f"максимально допустимо {config.MAX_PAGES}")
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = '\n'.join(text_parts)
            
            # Определяем кодировку
            detected_encoding = chardet.detect(full_text.encode())
            metadata['encoding'] = detected_encoding.get('encoding', 'utf-8')
            
            # Пробуем декодировать текст с определенной кодировкой
            full_text = self._try_decode(full_text, metadata['encoding'])
            
            return full_text, metadata
            
        except Exception as e:
            # Резервный вариант с PyPDF2
            return self._parse_pdf_fallback(file_path, metadata)
    
    def _parse_pdf_fallback(self, file_path: str, metadata: dict) -> Tuple[str, dict]:
        """Резервный метод парсинга PDF с использованием PyPDF2"""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata['pages'] = len(pdf_reader.pages)
            
            if metadata['pages'] > config.MAX_PAGES:
                raise ValueError(f"Документ содержит {metadata['pages']} страниц, "
                               f"максимально допустимо {config.MAX_PAGES}")
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        full_text = '\n'.join(text_parts)
        detected_encoding = chardet.detect(full_text.encode())
        metadata['encoding'] = detected_encoding.get('encoding', 'utf-8')
        
        return full_text, metadata
    
    def parse_docx(self, file_path: str) -> Tuple[str, dict]:
        """
        Парсит DOCX документ
        
        Args:
            file_path: путь к DOCX файлу
            
        Returns:
            Tuple[str, dict]: (текст документа, метаданные)
        """
        try:
            doc = Document(file_path)
            
            # Извлекаем текст из параграфов
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = '\n'.join(text_parts)
            
            # Определяем кодировку
            detected_encoding = chardet.detect(full_text.encode())
            encoding = detected_encoding.get('encoding', 'utf-8')
            
            # Пробуем декодировать текст
            full_text = self._try_decode(full_text, encoding)
            
            # Приблизительно оцениваем количество страниц
            # (примерно 400 слов на страницу)
            words = full_text.split()
            estimated_pages = max(1, len(words) // 400)
            
            metadata = {
                'pages': estimated_pages,
                'format': 'docx',
                'encoding': encoding
            }
            
            if metadata['pages'] > config.MAX_PAGES:
                raise ValueError(f"Документ содержит примерно {metadata['pages']} страниц, "
                               f"максимально допустимо {config.MAX_PAGES}")
            
            return full_text, metadata
            
        except Exception as e:
            raise Exception(f"Ошибка при парсинге DOCX: {str(e)}")
    
    def _try_decode(self, text: str, detected_encoding: str) -> str:
        """
        Пытается декодировать текст с различными кодировками
        
        Args:
            text: текст для декодирования
            detected_encoding: определенная кодировка
            
        Returns:
            str: декодированный текст
        """
        if detected_encoding and detected_encoding.lower() in [enc.lower() for enc in self.supported_encodings]:
            return text
        
        # Пробуем различные кодировки
        for encoding in self.supported_encodings:
            try:
                # Пробуем перекодировать
                encoded = text.encode('latin1', errors='ignore')
                decoded = encoded.decode(encoding, errors='ignore')
                if decoded and len(decoded) > len(text) * 0.8:  # Проверка на потерю данных
                    return decoded
            except:
                continue
        
        # Возвращаем оригинальный текст, если не удалось декодировать
        return text
    
    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Валидирует файл перед обработкой
        
        Args:
            file_path: путь к файлу
            
        Returns:
            Tuple[bool, Optional[str]]: (валидный ли файл, сообщение об ошибке)
        """
        # Проверяем существование файла
        if not os.path.exists(file_path):
            return False, "Файл не найден"
        
        # Проверяем расширение
        extension = os.path.splitext(file_path)[1].lower()
        if extension not in self.supported_formats:
            return False, f"Неподдерживаемый формат. Поддерживаются: {', '.join(self.supported_formats)}"
        
        # Проверяем размер файла
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > config.MAX_FILE_SIZE_MB:
            return False, f"Файл слишком большой ({file_size_mb:.1f} МБ). Максимум: {config.MAX_FILE_SIZE_MB} МБ"
        
        return True, None

